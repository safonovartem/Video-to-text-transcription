import os
import subprocess
import tempfile

from faster_whisper import WhisperModel

from docx import Document

from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer
)

from reportlab.lib.styles import getSampleStyleSheet


# ==========================================
# НАСТРОЙКИ
# ==========================================

VIDEO_PATH = r"C:\Users\PC\PycharmProjects\From_Audio_To_Text\input\2026-03-14 09-02-09.mp4"

TXT_OUTPUT = "transcript.txt"
DOCX_OUTPUT = "transcript.docx"
PDF_OUTPUT = "transcript.pdf"


# ==========================================
# ФОРМАТ ВРЕМЕНИ
# ==========================================

def format_time(seconds):

    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)

    return f"{hours:02}:{minutes:02}:{secs:02}"


# ==========================================
# ИЗВЛЕЧЕНИЕ АУДИО
# ==========================================

def extract_audio(video_path):

    temp_audio = tempfile.NamedTemporaryFile(
        suffix=".wav",
        delete=False
    )

    audio_path = temp_audio.name

    temp_audio.close()

    command = [
        "ffmpeg",
        "-i", video_path,
        "-vn",
        "-acodec", "pcm_s16le",
        "-ar", "16000",
        "-ac", "1",
        audio_path,
        "-y"
    ]

    subprocess.run(
        command,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL
    )

    return audio_path


# ==========================================
# ТРАНСКРИБАЦИЯ
# ==========================================

def transcribe(audio_path):

    model = WhisperModel(
        "small",
        device="cpu",
        compute_type="int8"
    )

    segments, info = model.transcribe(
        audio_path,
        language="ru",
        beam_size=1,
        vad_filter=True
    )

    transcript = []

    for segment in segments:

        start = format_time(segment.start)
        end = format_time(segment.end)

        text = segment.text.strip()

        line = f"[{start} - {end}] {text}"

        transcript.append(line)

        print(line)

    return transcript


# ==========================================
# TXT
# ==========================================

def save_txt(transcript):

    with open(TXT_OUTPUT, "w", encoding="utf-8") as file:

        for line in transcript:
            file.write(line + "\n")


# ==========================================
# DOCX
# ==========================================

def save_docx(transcript):

    document = Document()

    document.add_heading("Транскрибация", level=1)

    for line in transcript:
        document.add_paragraph(line)

    document.save(DOCX_OUTPUT)


# ==========================================
# PDF
# ==========================================

def save_pdf(transcript):

    doc = SimpleDocTemplate(PDF_OUTPUT)

    styles = getSampleStyleSheet()

    content = []

    for line in transcript:

        p = Paragraph(line, styles["BodyText"])

        content.append(p)

        content.append(Spacer(1, 5))

    doc.build(content)


# ==========================================
# MAIN
# ==========================================

def main():

    print("Извлечение аудио...")

    audio_path = extract_audio(VIDEO_PATH)

    print("Транскрибация...")

    transcript = transcribe(audio_path)

    print("Сохранение TXT...")
    save_txt(transcript)

    print("Сохранение DOCX...")
    save_docx(transcript)

    print("Сохранение PDF...")
    save_pdf(transcript)

    print("Удаление временного аудио...")

    os.remove(audio_path)

    print("Готово!")


if __name__ == "__main__":
    main()